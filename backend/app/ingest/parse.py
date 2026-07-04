"""Multimodal document parsing → normalized pages.

Supports PDF (text + scanned/OCR), Word (.docx, incl. its tables), and plain
text. Each parser returns a list of ``PageContent`` so downstream chunking is
format-agnostic.
"""

from __future__ import annotations

import logging
import os
from dataclasses import dataclass

from .ocr import ocr_image_bytes

logger = logging.getLogger(__name__)

# Below this many characters a PDF page is treated as scanned → OCR fallback.
_LOW_TEXT_THRESHOLD = 25


@dataclass
class PageContent:
    page: int
    text: str
    kind: str  # text | ocr


def parse_pdf(path: str) -> list[PageContent]:
    import fitz  # PyMuPDF

    pages: list[PageContent] = []
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            kind = "text"
            if len(text) < _LOW_TEXT_THRESHOLD:
                # Likely a scanned/image page — render and OCR it.
                pix = page.get_pixmap(dpi=200)
                ocr_text = ocr_image_bytes(pix.tobytes("png"))
                if ocr_text:
                    text, kind = ocr_text, "ocr"
            pages.append(PageContent(page=i, text=text, kind=kind))
    return pages


def parse_docx(path: str) -> list[PageContent]:
    import docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    # Word has no fixed pagination; treat the body as a single logical page.
    return [PageContent(page=1, text="\n".join(parts), kind="text")]


def parse_txt(path: str) -> list[PageContent]:
    with open(path, "r", encoding="utf-8", errors="ignore") as fh:
        return [PageContent(page=1, text=fh.read(), kind="text")]


def parse_document(path: str, mime: str) -> list[PageContent]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf" or "pdf" in mime:
        return parse_pdf(path)
    if ext in (".docx", ".doc") or "word" in mime or "officedocument" in mime:
        return parse_docx(path)
    return parse_txt(path)
